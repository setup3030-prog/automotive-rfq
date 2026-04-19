"""
Word (.docx) quote export service.

Generates a professional A4 customer-facing quote document.
Uses python-docx + BytesIO — no temp files (Vercel-safe).
"""
from __future__ import annotations

from datetime import datetime, timezone
from io import BytesIO

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ─── Colour palette (matches pdf_export.py) ───────────────────────────────────

_C_NAVY      = "193769"
_C_NAVY_MID  = "324B82"
_C_TOTAL_ROW = "BED2F5"
_C_LABEL     = "F2F4FA"
_C_WALK      = "FFDADA"
_C_AGG       = "FFFFD2"
_C_TARGET    = "DAFADA"
_C_GO        = "1A6E28"
_C_NOGO      = "AA1919"
_C_CAUTION   = "A07800"
_C_WHITE     = "FFFFFF"


# ─── Low-level XML helpers ────────────────────────────────────────────────────

def _shade_cell(cell, fill_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _shade_para(para, fill_hex: str) -> None:
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)


def _cell_text(cell, text: str, bold=False, size_pt=10,
               color_hex: str | None = None,
               align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = "Calibri"
    if color_hex:
        r, g, b = int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)


def _add_field_run(para, field_code: str) -> None:
    run = para.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = f" {field_code} "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(fld)
    run._r.append(instr)
    run._r.append(end)


def _remove_table_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tblBorders.append(el)
    tblPr.append(tblBorders)


# ─── Page setup ───────────────────────────────────────────────────────────────

def _set_page_margins(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)


def _set_header(doc: Document) -> None:
    section = doc.sections[0]
    section.different_first_page_header_footer = False
    header = section.header
    header.is_linked_to_previous = False
    header.paragraphs[0].clear()

    table = header.add_table(rows=1, cols=2, width=Cm(17.0))
    _remove_table_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.columns[0].width = Cm(11)
    table.columns[1].width = Cm(6)

    left_cell = table.cell(0, 0)
    p_left = left_cell.paragraphs[0]
    r = p_left.add_run("AUTOMOTIVE INJECTION MOLDING")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(9)
    r2 = p_left.add_run("  |  QUOTATION")
    r2.font.name = "Calibri"
    r2.font.size = Pt(9)

    right_cell = table.cell(0, 1)
    p_right = right_cell.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_str = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    r3 = p_right.add_run(f"{date_str}  |  CONFIDENTIAL")
    r3.font.name = "Calibri"
    r3.font.size = Pt(8)
    r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def _set_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.clear()

    r0 = para.add_run("Page ")
    r0.font.name = "Calibri"
    r0.font.size = Pt(8)
    r0.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    _add_field_run(para, "PAGE")
    r1 = para.add_run(" of ")
    r1.font.name = "Calibri"
    r1.font.size = Pt(8)
    r1.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    _add_field_run(para, "NUMPAGES")
    r2 = para.add_run("  |  Internal use — do not distribute")
    r2.font.name = "Calibri"
    r2.font.size = Pt(8)
    r2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


# ─── Document structure helpers ───────────────────────────────────────────────

def _section_heading(doc: Document, title: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(4)
    _shade_para(para, _C_NAVY)
    run = para.add_run(f"  {title.upper()}")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _label_value_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    _remove_table_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.columns[0].width = Cm(6)
    table.columns[1].width = Cm(11)

    for i, (label, value) in enumerate(rows):
        lc = table.cell(i, 0)
        vc = table.cell(i, 1)
        if i % 2 == 0:
            _shade_cell(lc, _C_LABEL)
        _cell_text(lc, label, bold=True, size_pt=9)
        _cell_text(vc, value, size_pt=9)


def _cost_table(doc: Document, rows: list[tuple[str, float]], total: float, cur: str) -> None:
    all_rows = rows + [("TOTAL", total)]
    table = doc.add_table(rows=len(all_rows) + 1, cols=3)
    _remove_table_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.columns[0].width = Cm(9.5)
    table.columns[1].width = Cm(4.5)
    table.columns[2].width = Cm(3.0)

    # Header row
    for col, txt in enumerate(["Component", f"Cost/Part ({cur})", "Share"]):
        _shade_cell(table.cell(0, col), _C_NAVY_MID)
        _cell_text(table.cell(0, col), txt, bold=True, size_pt=9, color_hex=_C_WHITE,
                   align=WD_ALIGN_PARAGRAPH.CENTER)

    # Body rows
    for i, (label, cost) in enumerate(rows):
        row_idx = i + 1
        if i % 2 == 0:
            for col in range(3):
                _shade_cell(table.cell(row_idx, col), _C_LABEL)
        share = f"{cost / total * 100:.1f}%" if total else "—"
        _cell_text(table.cell(row_idx, 0), label, size_pt=9)
        _cell_text(table.cell(row_idx, 1), f"{cost:.4f}", size_pt=9, align=WD_ALIGN_PARAGRAPH.RIGHT)
        _cell_text(table.cell(row_idx, 2), share, size_pt=9, align=WD_ALIGN_PARAGRAPH.RIGHT)

    # Total row
    t_idx = len(rows) + 1
    for col in range(3):
        _shade_cell(table.cell(t_idx, col), _C_TOTAL_ROW)
    _cell_text(table.cell(t_idx, 0), "TOTAL MANUFACTURING COST", bold=True, size_pt=9)
    _cell_text(table.cell(t_idx, 1), f"{total:.4f}", bold=True, size_pt=9, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _cell_text(table.cell(t_idx, 2), "100%", bold=True, size_pt=9, align=WD_ALIGN_PARAGRAPH.RIGHT)


def _pricing_table(doc: Document, rows: list[tuple[str, float, float, str]], cur: str) -> None:
    table = doc.add_table(rows=len(rows) + 1, cols=3)
    _remove_table_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.columns[0].width = Cm(6.0)
    table.columns[1].width = Cm(5.0)
    table.columns[2].width = Cm(6.0)

    for col, txt in enumerate(["Strategy", f"Unit Price ({cur})", "Gross Margin"]):
        _shade_cell(table.cell(0, col), _C_NAVY_MID)
        _cell_text(table.cell(0, col), txt, bold=True, size_pt=9, color_hex=_C_WHITE,
                   align=WD_ALIGN_PARAGRAPH.CENTER)

    fill_map = {"walk": _C_WALK, "agg": _C_AGG, "target": _C_TARGET}
    for i, (label, price, margin, fill_key) in enumerate(rows):
        ri = i + 1
        fill = fill_map.get(fill_key, _C_LABEL)
        for col in range(3):
            _shade_cell(table.cell(ri, col), fill)
        is_target = fill_key == "target"
        _cell_text(table.cell(ri, 0), label, bold=is_target, size_pt=9)
        _cell_text(table.cell(ri, 1), f"{price:.4f}", bold=is_target, size_pt=9, align=WD_ALIGN_PARAGRAPH.RIGHT)
        _cell_text(table.cell(ri, 2), f"{margin * 100:.1f}%", bold=is_target, size_pt=9, align=WD_ALIGN_PARAGRAPH.RIGHT)


def _decision_banner(doc: Document, text: str, fill_hex: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    _remove_table_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.columns[0].width = Cm(17.0)
    cell = table.cell(0, 0)
    _shade_cell(cell, fill_hex)
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


# ─── Public API ───────────────────────────────────────────────────────────────

def generate_quote_docx(data: dict) -> bytes:
    """
    Build a professional A4 .docx customer quote and return raw bytes.
    Accepts the same dict as generate_quote_pdf (PDFExportRequest.model_dump()).
    """
    doc = Document()
    _set_page_margins(doc)
    _set_header(doc)
    _set_footer(doc)

    cur   = data.get("currency", "PLN")
    total = float(data.get("total_cost") or 1)
    vol   = int(data.get("annual_volume") or 0)

    # ── Document title ────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.paragraph_format.space_after = Pt(4)
    tr = title_para.add_run("QUOTATION  —  ")
    tr.bold = True
    tr.font.name = "Calibri"
    tr.font.size = Pt(16)
    tr2 = title_para.add_run(str(data.get("rfq_name") or "RFQ"))
    tr2.bold = True
    tr2.font.name = "Calibri"
    tr2.font.size = Pt(16)
    tr2.font.color.rgb = RGBColor(0x19, 0x37, 0x69)

    sub_para = doc.add_paragraph()
    sub_para.paragraph_format.space_after = Pt(8)
    sr = sub_para.add_run(
        f"{data.get('customer') or ''}  |  Part: {data.get('part_number') or ''}  |  "
        f"Engineer: {data.get('quoting_engineer') or ''}  |  Date: {data.get('rfq_date') or ''}"
    )
    sr.font.name = "Calibri"
    sr.font.size = Pt(9)
    sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # ── Section 1: RFQ Details ────────────────────────────────────────────────
    _section_heading(doc, "1. RFQ Details")
    _label_value_table(doc, [
        ("Customer",          str(data.get("customer") or "—")),
        ("Project / RFQ",     str(data.get("rfq_name") or "—")),
        ("Part Number",       str(data.get("part_number") or "—")),
        ("Part Description",  str(data.get("part_description") or "—")),
        ("Quoting Engineer",  str(data.get("quoting_engineer") or "—")),
        ("RFQ Date",          str(data.get("rfq_date") or "—")),
        ("Currency",          cur),
        ("Annual Volume",     f"{vol:,} pcs"),
    ])

    # ── Section 2: Process Parameters ────────────────────────────────────────
    _section_heading(doc, "2. Process Parameters")
    _label_value_table(doc, [
        ("Cycle Time",        f"{data.get('cycle_time_s', 0):.1f} s"),
        ("Cavities",          str(data.get("cavities", 1))),
        ("OEE",               f"{data.get('oee_pct', 0):.1f}%"),
        ("Scrap Rate",        f"{data.get('scrap_rate_pct', 0):.1f}%"),
    ])

    # ── Section 3: Cost Breakdown ─────────────────────────────────────────────
    _section_heading(doc, "3. Cost Breakdown  (per part)")
    cost_rows: list[tuple[str, float]] = []
    for label, key in [
        ("Machine",            "machine_cost"),
        ("Material",           "material_cost"),
        ("Tooling / Amort.",   "tooling_cost"),
        ("Labor",              "labor_cost"),
        ("Energy",             "energy_cost"),
        ("Overhead",           "overhead_cost"),
        ("Logistics & Pack.",  "logistics_packaging"),
    ]:
        v = float(data.get(key) or 0)
        if v:
            cost_rows.append((label, v))
    _cost_table(doc, cost_rows, total, cur)

    # ── Section 4: Pricing Strategy ───────────────────────────────────────────
    _section_heading(doc, "4. Pricing Strategy")
    pricing_rows = [
        ("Walk Away (minimum)",  float(data.get("walk_away_price") or 0),  float(data.get("walk_away_margin") or 0),  "walk"),
        ("Aggressive (competitive)", float(data.get("aggressive_price") or 0), float(data.get("aggressive_margin") or 0), "agg"),
        ("Target (recommended)", float(data.get("target_price_calc") or 0),  float(data.get("target_margin") or 0),   "target"),
    ]
    _pricing_table(doc, pricing_rows, cur)

    cust_target = data.get("customer_target_price")
    if cust_target:
        ct_para = doc.add_paragraph()
        ct_para.paragraph_format.space_before = Pt(4)
        r = ct_para.add_run(f"Customer target price: {float(cust_target):.4f} {cur}  ")
        r.font.name = "Calibri"
        r.font.size = Pt(9)
        gap = float(data.get("target_price_calc") or 0) - float(cust_target)
        gap_str = f"+{gap:.4f}" if gap >= 0 else f"{gap:.4f}"
        r2 = ct_para.add_run(f"(gap vs. target: {gap_str} {cur})")
        r2.font.name = "Calibri"
        r2.font.size = Pt(9)
        r2.font.color.rgb = RGBColor(0xAA, 0x19, 0x19) if gap < 0 else RGBColor(0x1A, 0x6E, 0x28)

    # ── Section 5: Final Unit Price ───────────────────────────────────────────
    _section_heading(doc, "5. Final Unit Price")
    price_para = doc.add_paragraph()
    price_para.paragraph_format.space_before = Pt(6)
    pr = price_para.add_run(f"{float(data.get('target_price_calc') or 0):.4f} {cur}")
    pr.bold = True
    pr.font.name = "Calibri"
    pr.font.size = Pt(18)
    pr.font.color.rgb = RGBColor(0x19, 0x37, 0x69)

    annual_rev = float(data.get("target_price_calc") or 0) * vol
    rev_para = doc.add_paragraph()
    rr = rev_para.add_run(
        f"Annual project value at {vol:,} pcs:  "
        f"{annual_rev:,.0f} {cur}"
    )
    rr.font.name = "Calibri"
    rr.font.size = Pt(10)

    # ── Section 6: GO / NO-GO ─────────────────────────────────────────────────
    _section_heading(doc, "6. GO / NO-GO Decision")
    raw_decision = str(data.get("decision") or "")
    if "NO GO" in raw_decision.upper() or "HIGH RISK" in raw_decision.upper():
        fill = _C_NOGO
        banner_text = "NO GO — HIGH RISK"
    elif "CAUTION" in raw_decision.upper():
        fill = _C_CAUTION
        banner_text = "PROCEED WITH CAUTION"
    else:
        fill = _C_GO
        banner_text = "GO — QUOTE AT TARGET PRICE"
    _decision_banner(doc, banner_text, fill)

    # Note paragraph
    note_para = doc.add_paragraph()
    note_para.paragraph_format.space_before = Pt(8)
    note_r = note_para.add_run(
        "This quotation is valid for 30 days from the RFQ date. Prices are subject to change "
        "based on final volume confirmation and tooling specification. All prices are "
        f"exclusive of VAT and stated in {cur}."
    )
    note_r.font.name = "Calibri"
    note_r.font.size = Pt(8)
    note_r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
